from pathlib import Path

from jacowvalidator.docutils.doc import create_upload_variables, AbstractNotFoundError

test_dir = Path(__file__).parent / 'data'


def test_title():
    from docx import Document

    doc = Document(test_dir / 'test1.docx')
    summary, reference_csv_details, title = create_upload_variables(doc,)

    summary_detail = summary['Title']
    assert summary_detail['ok'], "Title has issues but should have no issues"
    assert len(summary_detail['details']), "Title is more than one paragraph"

    details = summary_detail['details'][0]
    assert details['style_ok'], "Title has style issues but should have no style issues"
    assert details['case_ok'], "Title has case issues but should have no case issues"
    assert details['title_style_ok'], "Title has strict style issues but should have no strict style issues"


def test_author():
    from docx import Document

    doc = Document(test_dir / 'test1.docx')
    summary, reference_csv_details, title = create_upload_variables(doc)

    summary_detail = summary['Authors']
    assert summary_detail['ok'], "Author has issues but should have no issues"
    assert len(summary_detail['details']), "Author is more than one paragraph"

    details = summary_detail['details'][0]
    assert details['style_ok'], "Author has style issues but should have no style issues"
    assert details['title_style_ok'], "Author has strict style issues but should have no strict style issues"


def test_abstract():
    from docx import Document

    doc = Document(test_dir / 'test1.docx')
    summary, reference_csv_details, title = create_upload_variables(doc)

    summary_detail = summary['Abstract']
    assert summary_detail['ok'], "Abstract has issues but should have no issues"
    assert len(summary_detail['details']), "Abstract is more than one paragraph"

    details = summary_detail['details'][0]
    assert details['style_ok'], "Abstract has style issues but should have no style issues"
    assert details['title_style_ok'], "Abstract has strict style issues but should have no strict style issues"


def test_no_abstract():
    from docx import Document

    doc = Document(test_dir / 'correct_a4_margins.docx')
    summary, reference_csv_details, title = None, None, None
    try:
        summary, reference_csv_details, title = create_upload_variables(doc)
        assert False, 'AbstractNotFoundError not raised'
    except AbstractNotFoundError as err:
        assert str(err) == "Abstract header not found", 'AbstractNotFoundError found but not correct message'
        # these should be empty because Error raised before values returned
        assert summary is None, 'summary not empty'
        assert reference_csv_details is None, 'reference_csv_details not empty'


def test_headings():
    from docx import Document

    doc = Document(test_dir / 'test1.docx')
    summary, reference_csv_details, title = create_upload_variables(doc)

    summary_detail = summary['Headings']
    assert summary_detail['ok'], "Headings has issues but should have no issues"

    for item in summary_detail['details']:
        assert item['style_ok'], "Headings has style issues but should have no style issues"


def test_paragraphs():
    from docx import Document

    doc = Document(test_dir / 'test1.docx')
    summary, reference_csv_details, title = create_upload_variables(doc)

    summary_detail = summary['Paragraphs']
    assert summary_detail['ok'], "Paragraphs has issues but should have no issues"

    for item in summary_detail['details']:
        assert item['style_ok'], "Paragraphs has style issues but should have no style issues"
